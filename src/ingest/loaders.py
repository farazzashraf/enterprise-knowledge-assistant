"""
File loaders for PDF, DOCX, and TXT with page/position tracking.
"""

import os
import re
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Union

logger = logging.getLogger(__name__)

@dataclass
class DocumentPage:
    """A single page or logical section of a document"""
    content: str
    page_number: int
    source: str #filepath
    metadata: Optional[dict] = None

@dataclass
class LoadedDocument:
    """A fully loaded document with all pages"""
    source: str
    pages: List[DocumentPage]
    file_type: str

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.content for p in self.pages)

class BaseLoader(ABC):
    """Abstract base class for document loaders"""
    @abstractmethod
    def load(self, file_path: Union[str, Path]) -> LoadedDocument:
        pass

    def _validate_path(self, file_path: Union[str, Path]) -> Path:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        return path

class PDFLoader(BaseLoader):
    """Extract text from PDF with real page numbers using PyMuPDF.

    When ``extract_images=True``, images embedded in each page (diagrams,
    screenshots, or whole scanned pages stored as a single image) are run
    through OCR and their text is appended to the page content — so a scanned
    PDF that has no selectable text still becomes searchable.
    """

    def __init__(self, extract_images: bool = False):
        self.extract_images = extract_images
        self._ocr_engine = None  # lazily constructed; OCR init is expensive

    def _get_ocr_engine(self):
        """Lazily build the RapidOCR engine. Returns None if OCR is unavailable."""
        if self._ocr_engine is not None:
            return self._ocr_engine
        try:
            from rapidocr_onnxruntime import RapidOCR
        except ImportError:
            logger.warning(
                "extract_images=True but 'rapidocr-onnxruntime' is not installed; "
                "skipping OCR. Install it with: pip install rapidocr-onnxruntime"
            )
            return None
        self._ocr_engine = RapidOCR()
        return self._ocr_engine

    def _ocr_page_images(self, doc, page) -> str:
        """OCR every raster image on a page; return the concatenated text."""
        engine = self._get_ocr_engine()
        if engine is None:
            return ""

        import io
        import numpy as np
        from PIL import Image

        ocr_texts: List[str] = []
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                base = doc.extract_image(xref)
                pil = Image.open(io.BytesIO(base["image"])).convert("RGB")
                result, _ = engine(np.array(pil))
                if result:
                    ocr_texts.append("\n".join(line[1] for line in result))
            except Exception as e:
                logger.warning(f"OCR failed for image xref {xref} on page: {e}")
        return "\n".join(t for t in ocr_texts if t.strip())

    def load(self, file_path: Union[str, Path]) -> LoadedDocument:
        path = self._validate_path(file_path)

        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError(
                "PyMuPDF is required for PDF loading. "
            )

        pages = []
        doc = fitz.open(str(path))
        try:
            for i, page in enumerate(doc, start=1):
                text = page.get_text("text")
                text = re.sub(r'\n\s*\n', '\n\n', text.strip()) if text else ""

                # Pull any text living inside images on this page via OCR.
                ocr_text = self._ocr_page_images(doc, page) if self.extract_images else ""
                if ocr_text:
                    text = f"{text}\n\n{ocr_text}".strip() if text else ocr_text

                if text:
                    rect = page.rect
                    pages.append(
                        DocumentPage(
                            content=text,
                            page_number=i,
                            source=str(path),
                            metadata={
                                "width": rect.width,
                                "height": rect.height,
                                "chars": len(text),
                                "ocr": bool(ocr_text),
                            }
                        )
                    )
        finally:
            doc.close()

        logger.info(f"Loaded PDF '{path.name}': {len(pages)} pages")
        return LoadedDocument(
            source=str(path),
            pages=pages,
            file_type="pdf"
        )

class DOCXLoader(BaseLoader):
    """Extract text from DOCX with simulated page numbers."""
    
    def load(self, file_path: Union[str, Path]) -> LoadedDocument:
        path = self._validate_path(file_path)
        
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX loading. "
                "Install it with: pip install python-docx"
            )
        
        doc = Document(path)
        pages = []
        current_page_num = 1
        current_page_text = []
        
        # DOCX doesn't have true pages, so we approximate by section breaks
        # or use a character-based heuristic (~3000 chars per page as fallback)
        PAGE_CHAR_THRESHOLD = 3000
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check for explicit page break
            if para._element.xpath('.//w:br[@w:type="page"]'):
                if current_page_text:
                    pages.append(
                        DocumentPage(
                            content="\n".join(current_page_text),
                            page_number=current_page_num,
                            source=str(path),
                            metadata={"approximate": True}
                        )
                    )
                    current_page_num += 1
                    current_page_text = []
            
            current_page_text.append(text)
            
            # Fallback heuristic if no page breaks found
            if len("\n".join(current_page_text)) > PAGE_CHAR_THRESHOLD:
                pages.append(
                    DocumentPage(
                        content="\n".join(current_page_text),
                        page_number=current_page_num,
                        source=str(path),
                        metadata={"approximate": True}
                    )
                )
                current_page_num += 1
                current_page_text = []
        
        # Add remaining content
        if current_page_text:
            pages.append(
                DocumentPage(
                    content="\n".join(current_page_text),
                    page_number=current_page_num,
                    source=str(path),
                    metadata={"approximate": True}
                )
            )
        
        # If no pages were created (empty doc), create one empty page
        if not pages:
            pages.append(
                DocumentPage(
                    content="",
                    page_number=1,
                    source=str(path),
                    metadata={"empty": True}
                )
            )
        
        logger.info(f"Loaded DOCX '{path.name}': {len(pages)} pages (approximate)")
        return LoadedDocument(
            source=str(path),
            pages=pages,
            file_type="docx"
        )

class TXTLoader(BaseLoader):
    """Load plain text with line-based page estimation."""
    
    def __init__(self, lines_per_page: int = 50):
        self.lines_per_page = lines_per_page
    
    def load(self, file_path: Union[str, Path]) -> LoadedDocument:
        path = self._validate_path(file_path)
        
        with open(path, "r", encoding="utf-8") as f:
            lines = f.readlines()
        
        # Clean lines
        lines = [line.rstrip() for line in lines]
        
        pages = []
        current_lines = []
        page_num = 1
        
        for line in lines:
            current_lines.append(line)
            if len(current_lines) >= self.lines_per_page:
                pages.append(
                    DocumentPage(
                        content="\n".join(current_lines),
                        page_number=page_num,
                        source=str(path),
                        metadata={"lines_per_page": self.lines_per_page}
                    )
                )
                current_lines = []
                page_num += 1
        
        # Add remaining lines
        if current_lines:
            pages.append(
                DocumentPage(
                    content="\n".join(current_lines),
                    page_number=page_num,
                    source=str(path),
                    metadata={"lines_per_page": self.lines_per_page}
                )
            )
        
        logger.info(f"Loaded TXT '{path.name}': {len(pages)} pages, {len(lines)} lines")
        return LoadedDocument(
            source=str(path),
            pages=pages,
            file_type="txt"
        )

class SmartLoader(BaseLoader):
    """Auto-detects file type and dispatches to the correct loader."""
    
    LOADERS = {
        ".pdf": PDFLoader,
        ".docx": DOCXLoader,
        ".txt": TXTLoader,
    }
    
    def __init__(self, **loader_kwargs):
        self.loader_kwargs = loader_kwargs
        self._loaders = {}
    
    def _get_loader(self, ext: str) -> BaseLoader:
        ext = ext.lower()
        if ext not in self._loaders:
            loader_cls = self.LOADERS.get(ext)
            if not loader_cls:
                raise ValueError(f"Unsupported file type: {ext}")
            self._loaders[ext] = loader_cls(**self.loader_kwargs.get(ext, {}))
        return self._loaders[ext]
    
    def load(self, file_path: Union[str, Path]) -> LoadedDocument:
        path = self._validate_path(file_path)
        ext = path.suffix
        loader = self._get_loader(ext)
        return loader.load(path)
    
    def load_directory(
        self,
        dir_path: Union[str, Path],
        extensions: Optional[List[str]] = None
    ) -> List[LoadedDocument]:
        """Load all supported files from a directory."""
        dir_path = Path(dir_path)
        if extensions is None:
            extensions = list(self.LOADERS.keys())
        
        docs = []
        for ext in extensions:
            for file_path in dir_path.rglob(f"*{ext}"):
                try:
                    doc = self.load(file_path)
                    docs.append(doc)
                    logger.info(f"  ✓ {file_path.name} ({len(doc.pages)} pages)")
                except Exception as e:
                    logger.warning(f"  ✗ Failed to load {file_path}: {e}")

        logger.info(f"Directory load complete: {len(docs)} documents loaded")
        return docs

# --- Convenience functions ---

def load_file(file_path: Union[str, Path]) -> LoadedDocument:
    """One-shot load with auto-detection."""
    return SmartLoader().load(file_path)


def load_directory(dir_path: Union[str, Path]) -> List[LoadedDocument]:
    """Load all supported files from a directory."""
    return SmartLoader().load_directory(dir_path)       

    